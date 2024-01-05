# Import necessary libraries
import cv2
import pytesseract
from argostranslate import package, translate
from docx import Document

def extract_text_from_image(image_path):
    img = cv2.imread(image_path)
    text = pytesseract.image_to_string(img)
    return text

def translate_text(text, source, target):
    # Load installed language packages
    installed_languages = translate.load_installed_languages()
    
    # Find the source and target languages
    source_lang = next((lang for lang in installed_languages if lang.code == source), None)
    target_lang = next((lang for lang in installed_languages if lang.code == target), None)
    
    # Check if the languages are installed
    if source_lang is None or target_lang is None:
        print(f"Could not find installed language packages for {source} or {target}. Trying to install...")
        
        # Update the package index
        package.update_package_index()

        # Get the available packages
        available_packages = package.get_available_packages()

        # Identify the package to install
        package_to_install = next(filter(lambda x: x.from_code == source and x.to_code == target, available_packages), None)

        # Download and install the package
        if package_to_install is not None:
            package.install_from_path(package_to_install.download())
            print(f"Installed language package for {source} to {target}.")
        else:
            print(f"No available language package for {source} to {target}.")
            return None

        # Reload installed languages
        installed_languages = translate.load_installed_languages()
        source_lang = next((lang for lang in installed_languages if lang.code == source), None)
        target_lang = next((lang for lang in installed_languages if lang.code == target), None)

    # Get a translation model
    translation_model = source_lang.get_translation(target_lang)
    
    # Translate the text
    result = translation_model.translate(text)
    return result

def save_to_word_doc(input_text, translated_text, file_path):
    doc = Document()
    doc.add_heading('Translation', 0)

    doc.add_heading('Input Text:', level=1)
    doc.add_paragraph(input_text)

    doc.add_heading('Translated Text:', level=1)
    doc.add_paragraph(translated_text)

    doc.save(file_path)

def main():
    # Install the language package
    package.install_from_path(r'C:\Users\OEM\Desktop\translate-en_ar-1_0.argosmodel')

    image_path = input("Enter the path to the image: ")
    source = input("Enter the source language: ")
    target = input("Enter the target language: ")
    file_path = input("Enter the full path (including file name and .docx extension) to save the Word document: ")

    text = extract_text_from_image(image_path)
    print(f"Extracted text in {source}: {text}")

    translated_text = translate_text(text, source, target)
    print(f"Translated text in {target}: {translated_text}")

    save_to_word_doc(text, translated_text, file_path)

# Run the main function
main()
